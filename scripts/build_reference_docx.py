"""Build the editable Word reference template shipped with DeepResearch-CLI."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = (
    PROJECT_ROOT / "assets" / "report-templates" / "word" / "reference.docx"
)


def _font(style, latin: str, east_asia: str, size: float, *, bold=False, color=None):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def _paragraph_style(styles, name: str, *, base: str = "Normal"):
    try:
        return styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[base]
        return style


def _character_style(styles, name: str, *, base: str = "Default Paragraph Font"):
    try:
        return styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        style.base_style = styles[base]
        return style


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("DEEP RESEARCH  ·  ")
    label.font.name = "Aptos"
    label.font.size = Pt(8)
    label.font.color.rgb = RGBColor.from_string("667085")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _paragraph_border(paragraph, *, side: str, color: str, size: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)
    borders.append(border)


def _style_border(style, *, side: str, color: str, size: str) -> None:
    properties = style.element.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)
    borders.append(border)


def _table_style(styles) -> None:
    try:
        style = styles["Table"]
    except KeyError:
        style = styles.add_style("Table", WD_STYLE_TYPE.TABLE)
    _font(style, "Aptos", "PingFang SC", 8.2, color="243447")

    properties = style.element.find(qn("w:tblPr"))
    if properties is None:
        properties = OxmlElement("w:tblPr")
        style.element.append(properties)
    margins = OxmlElement("w:tblCellMar")
    for side, width in (("top", "80"), ("left", "90"), ("bottom", "80"), ("right", "90")):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), width)
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    properties.append(margins)

    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "D8E0E8")
        borders.append(element)
    properties.append(borders)

    first_row = OxmlElement("w:tblStylePr")
    first_row.set(qn("w:type"), "firstRow")
    row_run = OxmlElement("w:rPr")
    bold = OxmlElement("w:b")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FFFFFF")
    row_run.extend((bold, color))
    row_cell = OxmlElement("w:tcPr")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "17324D")
    row_cell.append(shading)
    first_row.extend((row_run, row_cell))
    style.element.append(first_row)

    band = OxmlElement("w:tblStylePr")
    band.set(qn("w:type"), "band1Horz")
    band_cell = OxmlElement("w:tcPr")
    band_shading = OxmlElement("w:shd")
    band_shading.set(qn("w:val"), "clear")
    band_shading.set(qn("w:fill"), "F3F7FB")
    band_cell.append(band_shading)
    band.append(band_cell)
    style.element.append(band)


def build(output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(19)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(9)
    section.footer_distance = Mm(9)
    section.different_first_page_header_footer = True

    header = section.header.paragraphs[0]
    header.text = "DEEP RESEARCH  ·  深度研究报告"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.runs[0]
    header_run.font.name = "Aptos"
    header_run.font.size = Pt(8)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor.from_string("667085")
    _paragraph_border(header, side="bottom", color="D8E0E8", size="4")
    _page_number(section.footer.paragraphs[0])

    styles = document.styles
    normal = styles["Normal"]
    _font(normal, "Aptos", "PingFang SC", 9.8, color="243447")
    normal.paragraph_format.line_spacing = 1.4
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    _font(title, "Aptos Display", "PingFang SC", 27, bold=True, color="17324D")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.keep_with_next = True
    _style_border(title, side="top", color="0E7490", size="28")

    subtitle = styles["Subtitle"]
    _font(subtitle, "Aptos", "PingFang SC", 13, color="667085")
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.space_after = Pt(12)

    date_style = _paragraph_style(styles, "Date")
    _font(date_style, "Aptos", "PingFang SC", 9, color="667085")
    date_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name, size, before, after in (
        ("Heading 1", 18, 22, 11),
        ("Heading 2", 13.5, 16, 8),
        ("Heading 3", 11.5, 12, 6),
        ("Heading 4", 10.5, 10, 5),
    ):
        style = styles[name]
        _font(style, "Aptos Display", "PingFang SC", size, bold=True, color="17324D")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.page_break_before = True
    _style_border(styles["Heading 1"], side="bottom", color="D8E0E8", size="5")
    _font(styles["Heading 2"], "Aptos Display", "PingFang SC", 13.5, bold=True, color="2563A8")

    toc_heading = _paragraph_style(styles, "TOC Heading", base="Heading 1")
    _font(toc_heading, "Aptos Display", "PingFang SC", 22, bold=True, color="17324D")
    toc_heading.paragraph_format.page_break_before = True
    toc_heading.paragraph_format.space_after = Pt(16)

    compact = _paragraph_style(styles, "Compact")
    _font(compact, "Aptos", "PingFang SC", 9.2, color="243447")
    compact.paragraph_format.line_spacing = 1.3
    compact.paragraph_format.space_after = Pt(3)
    compact.paragraph_format.first_line_indent = Mm(0)

    for name in ("Body Text", "First Paragraph"):
        body = _paragraph_style(styles, name)
        _font(body, "Aptos", "PingFang SC", 9.8, color="243447")
        body.paragraph_format.line_spacing = 1.4
        body.paragraph_format.space_before = Pt(0)
        body.paragraph_format.space_after = Pt(0)
        body.paragraph_format.widow_control = True
    styles["First Paragraph"].paragraph_format.first_line_indent = Mm(0)

    list_paragraph = _paragraph_style(styles, "List Paragraph")
    _font(list_paragraph, "Aptos", "PingFang SC", 9.8, color="243447")
    list_paragraph.paragraph_format.line_spacing = 1.4
    list_paragraph.paragraph_format.space_before = Pt(0)
    list_paragraph.paragraph_format.space_after = Pt(0)
    list_paragraph.paragraph_format.left_indent = Mm(7.4)
    list_paragraph.paragraph_format.first_line_indent = Mm(-3.7)

    _font(styles["Quote"], "Aptos", "PingFang SC", 9.5, color="475467")
    styles["Quote"].paragraph_format.left_indent = Mm(8)
    styles["Quote"].paragraph_format.right_indent = Mm(8)
    styles["Quote"].paragraph_format.space_before = Pt(0)
    styles["Quote"].paragraph_format.space_after = Pt(0)

    source_code = _paragraph_style(styles, "Source Code")
    _font(source_code, "JetBrains Mono", "PingFang SC", 8.5, color="404040")
    source_code.paragraph_format.line_spacing = 1.15
    source_code.paragraph_format.space_before = Pt(0)
    source_code.paragraph_format.space_after = Pt(0)
    source_code.paragraph_format.left_indent = Mm(5)
    source_code.paragraph_format.right_indent = Mm(5)

    verbatim = _character_style(styles, "Verbatim Char")
    _font(verbatim, "JetBrains Mono", "PingFang SC", 9, color="404040")
    inline_code = _character_style(styles, "Inline Code")
    _font(inline_code, "JetBrains Mono", "PingFang SC", 9, color="404040")

    caption = _paragraph_style(styles, "Caption")
    _font(caption, "Aptos", "PingFang SC", 8.5, color="667085")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)
    caption.paragraph_format.keep_with_next = True
    for name in ("Figure", "Image Caption", "Table Caption"):
        item = _paragraph_style(styles, name, base="Caption")
        _font(item, "Aptos", "PingFang SC", 8.5, color="667085")

    for level, size, indent in ((1, 13, 0), (2, 12, 5.5), (3, 11, 11)):
        toc = _paragraph_style(styles, f"TOC {level}")
        _font(toc, "Aptos", "PingFang SC", size, bold=level == 1, color="243447")
        toc.paragraph_format.left_indent = Mm(indent)
        toc.paragraph_format.space_before = Pt(1 if level == 1 else 0)
        toc.paragraph_format.space_after = Pt(1)

    _table_style(styles)

    document.core_properties.title = "DeepResearch 中文研究报告模板"
    document.core_properties.subject = "由 Markdown 生成的可编辑研究报告"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.output.expanduser().resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
